#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
تحويل ملفات الأمان من Markdown إلى PDF و Word
"""

import os
import sys
from pathlib import Path

def check_dependencies():
    """التحقق من المكتبات المطلوبة"""
    required = ['markdown', 'python-docx', 'markdown2']
    missing = []
    
    for package in required:
        try:
            if package == 'python-docx':
                __import__('docx')
            else:
                __import__(package.replace('-', '_'))
        except ImportError:
            missing.append(package)
    
    if missing:
        print(f"❌ المكتبات التالية مفقودة: {', '.join(missing)}")
        print(f"\nقم بتثبيتها باستخدام:")
        print(f"pip install {' '.join(missing)}")
        return False
    
    return True

def markdown_to_word(md_file, output_file):
    """تحويل Markdown إلى Word"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import markdown2
        
        # قراءة ملف Markdown
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # تحويل Markdown إلى HTML
        html = markdown2.markdown(md_content, extras=['tables', 'fenced-code-blocks'])
        
        # إنشاء مستند Word
        doc = Document()
        
        # إضافة عنوان
        title = doc.add_heading(Path(md_file).stem.replace('-', ' ').title(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # معالجة النص (بسيطة - يمكن تحسينها)
        lines = md_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif line.startswith('```'):
                continue
            else:
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # حفظ المستند
        doc.save(output_file)
        return True
        
    except Exception as e:
        print(f"❌ خطأ في تحويل {md_file}: {str(e)}")
        return False

def markdown_to_pdf_html(md_file, output_file):
    """تحويل Markdown إلى PDF عبر HTML"""
    try:
        import markdown2
        
        # قراءة ملف Markdown
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # تحويل إلى HTML
        html = markdown2.markdown(md_content, extras=['tables', 'fenced-code-blocks', 'code-friendly'])
        
        # إنشاء HTML كامل مع CSS
        full_html = f"""
<!DOCTYPE html>
<html dir="rtl" lang="ar">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{Path(md_file).stem}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            direction: rtl;
        }}
        h1, h2, h3 {{ color: #0d6a79; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 12px; text-align: right; }}
        th {{ background-color: #0d6a79; color: white; }}
        code {{ background: #f4f4f4; padding: 2px 5px; border-radius: 3px; }}
        pre {{ background: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }}
        blockquote {{ border-right: 4px solid #0d6a79; padding-right: 15px; color: #666; }}
        .checkmark {{ color: green; font-weight: bold; }}
    </style>
</head>
<body>
{html}
</body>
</html>
"""
        
        # حفظ HTML
        html_output = output_file.replace('.pdf', '.html')
        with open(html_output, 'w', encoding='utf-8') as f:
            f.write(full_html)
        
        print(f"✅ تم إنشاء HTML: {html_output}")
        print(f"💡 يمكنك فتح الملف في المتصفح وطباعته كـ PDF")
        
        return True
        
    except Exception as e:
        print(f"❌ خطأ في تحويل {md_file}: {str(e)}")
        return False

def main():
    """الدالة الرئيسية"""
    print("=" * 60)
    print("🔐 تحويل ملفات الأمان إلى PDF و Word")
    print("=" * 60)
    
    # التحقق من المكتبات
    if not check_dependencies():
        return 1
    
    # الملفات المطلوب تحويلها
    md_files = [
        'SECURITY-IMPLEMENTATION-PLAN.md',
        'SECURITY-COMPLETION-SUMMARY.md',
        'SECURITY-FINAL-REPORT.md'
    ]
    
    success_count = 0
    
    for md_file in md_files:
        if not os.path.exists(md_file):
            print(f"⚠️  الملف غير موجود: {md_file}")
            continue
        
        print(f"\n📄 معالجة: {md_file}")
        
        # تحويل إلى Word
        word_output = md_file.replace('.md', '.docx')
        if markdown_to_word(md_file, word_output):
            print(f"  ✅ Word: {word_output}")
            success_count += 1
        
        # تحويل إلى HTML (للطباعة كـ PDF)
        pdf_output = md_file.replace('.md', '.pdf')
        if markdown_to_pdf_html(md_file, pdf_output):
            success_count += 1
    
    print("\n" + "=" * 60)
    print(f"✅ تم التحويل بنجاح: {success_count} ملف")
    print("=" * 60)
    
    return 0

if __name__ == '__main__':
    sys.exit(main())
